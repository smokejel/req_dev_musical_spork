"""
Document parsing utilities for extracting text from various file formats.

Supports:
- Plain text files (.txt)
- Word documents (.docx)
- PDF files (.pdf)
"""

import os
from pathlib import Path
from typing import Tuple
import docx
import PyPDF2


class DocumentParseError(Exception):
    """Raised when document parsing fails."""
    pass


def parse_txt(file_path: str) -> str:
    """
    Parse a plain text file.

    Args:
        file_path: Path to the .txt file

    Returns:
        Extracted text content

    Raises:
        DocumentParseError: If file cannot be read
    """
    try:
        # Try UTF-8 first, fall back to latin-1
        encodings = ['utf-8', 'latin-1', 'cp1252']

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                return content
            except UnicodeDecodeError:
                continue

        # If all encodings fail
        raise DocumentParseError(
            f"Failed to decode {file_path} with encodings: {encodings}"
        )

    except FileNotFoundError:
        raise DocumentParseError(f"File not found: {file_path}")
    except Exception as e:
        raise DocumentParseError(f"Error reading text file: {str(e)}")


def parse_docx(file_path: str) -> str:
    """
    Parse a Word document (.docx).

    Args:
        file_path: Path to the .docx file

    Returns:
        Extracted text content

    Raises:
        DocumentParseError: If file cannot be parsed
    """
    try:
        doc = docx.Document(file_path)

        # Extract text from paragraphs
        paragraphs = [para.text for para in doc.paragraphs]

        # Extract text from tables
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                table_text.append(" | ".join(row_text))

        # Combine all text
        all_text = paragraphs + table_text
        content = "\n".join(all_text)

        return content

    except FileNotFoundError:
        raise DocumentParseError(f"File not found: {file_path}")
    except Exception as e:
        raise DocumentParseError(f"Error parsing DOCX file: {str(e)}")


def parse_pdf(file_path: str) -> str:
    """
    Parse a PDF document.

    Args:
        file_path: Path to the .pdf file

    Returns:
        Extracted text content

    Raises:
        DocumentParseError: If file cannot be parsed
    """
    try:
        content_parts = []

        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)

            # Extract text from each page
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()

                if text.strip():  # Only add non-empty pages
                    content_parts.append(f"--- Page {page_num + 1} ---")
                    content_parts.append(text)

        content = "\n".join(content_parts)

        if not content.strip():
            raise DocumentParseError(
                f"No text could be extracted from PDF: {file_path}"
            )

        return content

    except FileNotFoundError:
        raise DocumentParseError(f"File not found: {file_path}")
    except Exception as e:
        raise DocumentParseError(f"Error parsing PDF file: {str(e)}")


def parse_document(file_path: str) -> Tuple[str, str]:
    """
    Parse a document and return its text content.

    Automatically detects file type based on extension and uses
    the appropriate parser.

    Args:
        file_path: Path to the document file

    Returns:
        Tuple of (content, file_type) where file_type is 'txt', 'docx', or 'pdf'

    Raises:
        DocumentParseError: If file type is unsupported or parsing fails
    """
    # Validate file exists
    path = Path(file_path)
    if not path.exists():
        raise DocumentParseError(f"File not found: {file_path}")

    if not path.is_file():
        raise DocumentParseError(f"Not a file: {file_path}")

    # Get file extension
    extension = path.suffix.lower()

    # Dispatch to appropriate parser
    if extension == '.txt':
        content = parse_txt(file_path)
        return content, 'txt'

    elif extension == '.docx':
        content = parse_docx(file_path)
        return content, 'docx'

    elif extension == '.pdf':
        content = parse_pdf(file_path)
        return content, 'pdf'

    else:
        raise DocumentParseError(
            f"Unsupported file type: {extension}. "
            f"Supported types: .txt, .docx, .pdf"
        )


def validate_document_content(content: str, min_length: int = 10) -> bool:
    """
    Validate that document content is non-empty and meaningful.

    Args:
        content: Document text content
        min_length: Minimum acceptable length in characters

    Returns:
        True if valid, False otherwise
    """
    if not content:
        return False

    # Remove whitespace for validation
    stripped = content.strip()

    if len(stripped) < min_length:
        return False

    return True
