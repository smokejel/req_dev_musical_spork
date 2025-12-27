"""
Export Service.

Generates exports in multiple formats: Markdown, DOCX, CSV, JSON, ZIP.
"""

import csv
import json
import io
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, List, Optional

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from api.models.database import WorkflowRun
from src.state import DecompositionState


class ExportService:
    """Service for exporting workflow results in various formats."""

    @staticmethod
    def export_markdown(workflow: WorkflowRun, state: DecompositionState) -> str:
        """
        Export workflow results as Markdown.

        Args:
            workflow: WorkflowRun database record
            state: Final state from checkpoint

        Returns:
            Markdown-formatted string
        """
        lines = []

        # Header
        lines.append(f"# Requirements Decomposition Report")
        lines.append(f"")
        lines.append(f"**Project:** {workflow.project_name}")
        lines.append(f"**Source Document:** {workflow.source_document}")
        lines.append(f"**Generated:** {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')} UTC")
        lines.append(f"**Status:** {workflow.status.value.upper()}")
        lines.append(f"")

        # Configuration
        config = workflow.config or {}
        lines.append(f"## Configuration")
        lines.append(f"")
        lines.append(f"- **Target Subsystem:** {config.get('subsystem', 'N/A')}")
        lines.append(f"- **Domain:** {config.get('domain', 'generic')}")
        lines.append(f"- **Quality Threshold:** {config.get('quality_threshold', 0.80)}")
        lines.append(f"- **Max Iterations:** {config.get('max_iterations', 3)}")
        lines.append(f"- **Review Mode:** {config.get('review_mode', 'before')}")
        lines.append(f"")

        # Summary
        lines.append(f"## Summary")
        lines.append(f"")
        lines.append(f"- **Extracted Requirements:** {workflow.extracted_count or 0}")
        lines.append(f"- **Generated Requirements:** {workflow.generated_count or 0}")
        lines.append(f"- **Quality Score:** {workflow.quality_score or 0:.2f}")
        lines.append(f"- **Elapsed Time:** {workflow.elapsed_time or 0:.2f}s")
        lines.append(f"- **Total Cost:** ${workflow.total_cost or 0:.4f}")
        lines.append(f"- **Energy:** {workflow.energy_wh or 0:.4f} Wh")
        lines.append(f"")

        # Decomposed Requirements
        decomposed = state.get("decomposed_requirements", [])
        if decomposed:
            lines.append(f"## Decomposed Requirements")
            lines.append(f"")

            for req in decomposed:
                lines.append(f"### {req.get('id', 'UNKNOWN')}")
                lines.append(f"")
                lines.append(f"**Text:** {req.get('text', 'N/A')}")
                lines.append(f"")

                if req.get("rationale"):
                    lines.append(f"**Rationale:** {req['rationale']}")
                    lines.append(f"")

                acceptance = req.get("acceptance_criteria", [])
                if acceptance:
                    lines.append(f"**Acceptance Criteria:**")
                    lines.append(f"")
                    for criterion in acceptance:
                        lines.append(f"- {criterion}")
                    lines.append(f"")

                lines.append(f"**Category:** {req.get('category', 'N/A')}")
                lines.append(f"**Priority:** {req.get('priority', 'N/A')}")

                if req.get("parent_id"):
                    lines.append(f"**Parent:** {req['parent_id']}")

                lines.append(f"")
                lines.append(f"---")
                lines.append(f"")

        # Quality Metrics
        quality = state.get("quality_metrics", {})
        if quality:
            lines.append(f"## Quality Metrics")
            lines.append(f"")
            lines.append(f"- **Overall Score:** {quality.get('overall_score', 0):.2f}")
            lines.append(f"- **Completeness:** {quality.get('completeness', 0):.2f}")
            lines.append(f"- **Clarity:** {quality.get('clarity', 0):.2f}")
            lines.append(f"- **Testability:** {quality.get('testability', 0):.2f}")
            lines.append(f"- **Traceability:** {quality.get('traceability', 0):.2f}")

            if quality.get("domain_compliance"):
                lines.append(f"- **Domain Compliance:** {quality['domain_compliance']:.2f}")

            lines.append(f"")

        # Validation Issues
        issues = state.get("validation_issues", [])
        if issues:
            lines.append(f"## Validation Issues")
            lines.append(f"")
            for issue in issues:
                lines.append(f"- {issue}")
            lines.append(f"")

        return "\n".join(lines)

    @staticmethod
    def export_docx(workflow: WorkflowRun, state: DecompositionState) -> bytes:
        """
        Export workflow results as Word document.

        Args:
            workflow: WorkflowRun database record
            state: Final state from checkpoint

        Returns:
            DOCX file as bytes
        """
        doc = Document()

        # Title
        title = doc.add_heading("Requirements Decomposition Report", 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Metadata
        doc.add_paragraph(f"Project: {workflow.project_name}")
        doc.add_paragraph(f"Source: {workflow.source_document}")
        doc.add_paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')} UTC")
        doc.add_paragraph(f"Status: {workflow.status.value.upper()}")
        doc.add_paragraph("")

        # Configuration
        doc.add_heading("Configuration", 1)
        config = workflow.config or {}
        doc.add_paragraph(f"Target Subsystem: {config.get('subsystem', 'N/A')}")
        doc.add_paragraph(f"Domain: {config.get('domain', 'generic')}")
        doc.add_paragraph(f"Quality Threshold: {config.get('quality_threshold', 0.80)}")
        doc.add_paragraph(f"Max Iterations: {config.get('max_iterations', 3)}")
        doc.add_paragraph("")

        # Summary
        doc.add_heading("Summary", 1)
        doc.add_paragraph(f"Extracted Requirements: {workflow.extracted_count or 0}")
        doc.add_paragraph(f"Generated Requirements: {workflow.generated_count or 0}")
        doc.add_paragraph(f"Quality Score: {workflow.quality_score or 0:.2f}")
        doc.add_paragraph(f"Elapsed Time: {workflow.elapsed_time or 0:.2f}s")
        doc.add_paragraph(f"Total Cost: ${workflow.total_cost or 0:.4f}")
        doc.add_paragraph("")

        # Requirements
        decomposed = state.get("decomposed_requirements", [])
        if decomposed:
            doc.add_heading("Decomposed Requirements", 1)

            for req in decomposed:
                doc.add_heading(req.get("id", "UNKNOWN"), 2)
                doc.add_paragraph(f"Text: {req.get('text', 'N/A')}")

                if req.get("rationale"):
                    doc.add_paragraph(f"Rationale: {req['rationale']}")

                acceptance = req.get("acceptance_criteria", [])
                if acceptance:
                    p = doc.add_paragraph("Acceptance Criteria:")
                    for criterion in acceptance:
                        doc.add_paragraph(criterion, style='List Bullet')

                doc.add_paragraph(f"Category: {req.get('category', 'N/A')}")
                doc.add_paragraph(f"Priority: {req.get('priority', 'N/A')}")
                doc.add_paragraph("")

        # Quality Metrics
        quality = state.get("quality_metrics", {})
        if quality:
            doc.add_heading("Quality Metrics", 1)
            doc.add_paragraph(f"Overall Score: {quality.get('overall_score', 0):.2f}")
            doc.add_paragraph(f"Completeness: {quality.get('completeness', 0):.2f}")
            doc.add_paragraph(f"Clarity: {quality.get('clarity', 0):.2f}")
            doc.add_paragraph(f"Testability: {quality.get('testability', 0):.2f}")
            doc.add_paragraph(f"Traceability: {quality.get('traceability', 0):.2f}")

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    @staticmethod
    def export_csv(workflow: WorkflowRun, state: DecompositionState) -> str:
        """
        Export requirements as CSV.

        Args:
            workflow: WorkflowRun database record
            state: Final state from checkpoint

        Returns:
            CSV-formatted string
        """
        output = io.StringIO()
        writer = csv.DictWriter(
            output,
            fieldnames=[
                "id",
                "text",
                "category",
                "priority",
                "parent_id",
                "rationale",
                "acceptance_criteria"
            ]
        )

        writer.writeheader()

        decomposed = state.get("decomposed_requirements", [])
        for req in decomposed:
            writer.writerow({
                "id": req.get("id", ""),
                "text": req.get("text", ""),
                "category": req.get("category", ""),
                "priority": req.get("priority", ""),
                "parent_id": req.get("parent_id", ""),
                "rationale": req.get("rationale", ""),
                "acceptance_criteria": "; ".join(req.get("acceptance_criteria", []))
            })

        return output.getvalue()

    @staticmethod
    def export_json(workflow: WorkflowRun, state: DecompositionState) -> str:
        """
        Export full state as JSON.

        Args:
            workflow: WorkflowRun database record
            state: Final state from checkpoint

        Returns:
            JSON-formatted string
        """
        export_data = {
            "workflow": {
                "id": workflow.id,
                "project_name": workflow.project_name,
                "source_document": workflow.source_document,
                "status": workflow.status.value,
                "created_at": workflow.created_at.isoformat() if workflow.created_at else None,
                "completed_at": workflow.completed_at.isoformat() if workflow.completed_at else None,
                "quality_score": workflow.quality_score,
                "total_cost": workflow.total_cost,
                "elapsed_time": workflow.elapsed_time,
                "config": workflow.config
            },
            "state": {
                "extracted_requirements": state.get("extracted_requirements", []),
                "decomposed_requirements": state.get("decomposed_requirements", []),
                "quality_metrics": state.get("quality_metrics", {}),
                "validation_issues": state.get("validation_issues", []),
                "traceability_matrix": state.get("traceability_matrix", {}),
                "system_context": state.get("system_context"),
                "decomposition_strategy": state.get("decomposition_strategy"),
                "cost_breakdown": state.get("cost_breakdown", {}),
                "timing_breakdown": state.get("timing_breakdown", {}),
                "energy_breakdown": state.get("energy_breakdown", {}),
                "token_usage": state.get("token_usage", {})
            }
        }

        return json.dumps(export_data, indent=2, default=str)

    @staticmethod
    def export_zip(workflow: WorkflowRun, state: DecompositionState) -> bytes:
        """
        Export all formats bundled in a ZIP file.

        Args:
            workflow: WorkflowRun database record
            state: Final state from checkpoint

        Returns:
            ZIP file as bytes
        """
        buffer = io.BytesIO()

        with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
            # Add Markdown
            md_content = ExportService.export_markdown(workflow, state)
            zf.writestr(f"{workflow.project_name}_report.md", md_content)

            # Add DOCX
            docx_content = ExportService.export_docx(workflow, state)
            zf.writestr(f"{workflow.project_name}_report.docx", docx_content)

            # Add CSV
            csv_content = ExportService.export_csv(workflow, state)
            zf.writestr(f"{workflow.project_name}_requirements.csv", csv_content)

            # Add JSON
            json_content = ExportService.export_json(workflow, state)
            zf.writestr(f"{workflow.project_name}_full_data.json", json_content)

        buffer.seek(0)
        return buffer.read()
